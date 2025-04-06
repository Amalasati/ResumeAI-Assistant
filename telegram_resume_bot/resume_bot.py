import telebot
from docx import Document
from docx2pdf import convert
import os
from dotenv import load_dotenv

load_dotenv()

BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")  
bot = telebot.TeleBot(BOT_TOKEN)

@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "Hi! Send me a skill to add to your resume (e.g., AWS, Azure, Docker).")

@bot.message_handler(func=lambda msg: True)
def handle_skill(message):
    skills = message.text.strip().split(',')
    skills = [skill.strip() for skill in skills if skill.strip()]
    
    if not skills:
        bot.reply_to(message, "Please provide at least one skill.")
        return
    
    try:
        resume_path = r"telegram_resume_bot\Amal_Asati__Resume.docx"
        doc = Document(resume_path)
        
        # Find the Skills section
        skills_para = None
        skills_index = -1
        
        for i, para in enumerate(doc.paragraphs):
            if "SKILLS" in para.text:
                skills_para = para
                skills_index = i
                break
        
        if skills_para is None:
            bot.reply_to(message, "Could not find the 'SKILLS' section.")
            return
        
        # Get formatting from a template paragraph
        template_para = None
        for i in range(skills_index + 1, len(doc.paragraphs)):
            if "•" in doc.paragraphs[i].text or "|" in doc.paragraphs[i].text:
                template_para = doc.paragraphs[i]
                break
        
        # Find the target bullet point to append to
        target_para = None
        for para in doc.paragraphs[skills_index + 1:]:
            if "MySQL | PowerBI | MS-Excel" in para.text:
                target_para = para
                break

        if not target_para:
            bot.reply_to(message, "Could not find the target skills line to update.")
            return

        # Append skills to the existing text
        existing_text = target_para.text.strip()
        if not existing_text.endswith('|'):
            existing_text += ' | '
        existing_text += ' | '.join(skills)

        # Clear and update paragraph
        target_para.clear()
        run = target_para.add_run(existing_text)

        if template_para and template_para.runs:
            run.font.size = template_para.runs[0].font.size
            run.font.name = template_para.runs[0].font.name
            run.font.bold = template_para.runs[0].font.bold
            run.font.italic = template_para.runs[0].font.italic
        if template_para:
            target_para.paragraph_format.left_indent = template_para.paragraph_format.left_indent
            target_para.paragraph_format.line_spacing = template_para.paragraph_format.line_spacing
            target_para.paragraph_format.space_after = template_para.paragraph_format.space_after
            target_para.paragraph_format.space_before = template_para.paragraph_format.space_before
        
        # Save and convert
        updated_docx = "updated_resume.docx"
        updated_pdf = "updated_resume.pdf"
        doc.save(updated_docx)
        convert(updated_docx, updated_pdf)
        
        with open(updated_pdf, "rb") as f:
            bot.send_document(message.chat.id, f)

        os.remove(updated_docx)
        os.remove(updated_pdf)
        
        bot.reply_to(message, f"✅ Successfully added: {', '.join(skills)}")

    except Exception as e:
        bot.reply_to(message, f"❌ An error occurred: {e}")

print("Bot is running...")
bot.infinity_polling()
