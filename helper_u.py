import google.generativeai as genai
import PyPDF2 as pdf
import json
import io
from docx import Document
from docx2pdf import convert
import tempfile
import os

def configure_genai(api_key):
    """Configure the Generative AI API with error handling."""
    try:
        genai.configure(api_key=api_key)
    except Exception as e:
        raise Exception(f"Failed to configure Generative AI: {str(e)}")
    

def get_gemini_response(prompt):
    """Generate a response using Gemini with enhanced error handling and response validation."""
    try:
        model = genai.GenerativeModel('gemini-1.5-pro')
        response = model.generate_content(prompt)
        
        # Ensure response is not empty
        if not response or not response.text:
            raise Exception("Empty response received from Gemini")
            
        # Try to parse the response as JSON
        try:
            response_json = json.loads(response.text)
            
            # Validate required fields
            required_fields = ["JD Match", "MissingKeywords", "Profile Summary", "Improvements"]
            for field in required_fields:
                if field not in response_json:
                    raise ValueError(f"Missing required field: {field}")
                    
            return response.text
            
        except json.JSONDecodeError:
            # If response is not valid JSON, try to extract JSON-like content
            import re
            json_pattern = r'\{.*\}'
            match = re.search(json_pattern, response.text, re.DOTALL)
            if match:
                return match.group()
            else:
                raise Exception("Could not extract valid JSON response")
                
    except Exception as e:
        raise Exception(f"Error generating response: {str(e)}")

def extract_pdf_text(uploaded_file):
    """Extract text from PDF with enhanced error handling."""
    try:
        reader = pdf.PdfReader(uploaded_file)
        if len(reader.pages) == 0:
            raise Exception("PDF file is empty")
            
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
                
        if not text:
            raise Exception("No text could be extracted from the PDF")
            
        return " ".join(text)
        
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")
    
def extract_docx_text(docx_file):
    """Extract text from Word document with error handling."""
    try:
        doc = Document(docx_file)
        
        # Extract text from paragraphs
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        if not text:
            raise Exception("No text could be extracted from the Word document")
            
        return " ".join(text)
        
    except Exception as e:
        raise Exception(f"Error extracting Word document text: {str(e)}")

def prepare_prompt(resume_text, job_description):
    """Prepare the input prompt with improved structure and validation."""
    if not resume_text or not job_description:
        raise ValueError("Resume text and job description cannot be empty")
        
    prompt_template = """
    Act as an expert Applicant Tracking System (ATS) specialist with deep understanding in:
    - Technical fields
    - Software engineering
    - Data science
    - Data analysis
    - Big data engineering
    
    Evaluate the following resume against the job description. Consider that the job market is highly competitive. Provide detailed feedback for resume improvement.
    
    Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Based on the evaluation, generate a tailored cold message to send to a recruiter or hiring manager for the position, highlighting the key qualifications and why the candidate is a strong fit for the role. Limit the cold message to 150-200 words.
    
    Also, provide specific improvements that can be added to the resume to better match the job description. These improvements should be categorized into:
    1. Experience - bullet points that could be added to existing experience sections
    2. Skills - specific skills that should be added or highlighted
    
    Provide a response in the following JSON format ONLY:
    {{
        "JD Match": A percentage (0-100) indicating the overall alignment of the resume with the job description,
        "MissingKeywords": List of specific skills, tools, or phrases missing in the resume but present in the job description,
        "Profile Summary": Detailed analysis of the match and specific improvement suggestions,
        "Improvements": {{
            "Experience": [List of experience bullet points to add],
            "Skills": [List of skills to add or highlight],
            "Projects": [List of project descriptions or achievements]
        }},
        "Cold Message": A persuasive, tailored message (50-100 words) highlighting key qualifications and enthusiasm for the role, designed to attract the recruiter or hiring manager
    }}
    """
    
    return prompt_template.format(
        resume_text=resume_text.strip(),
        job_description=job_description.strip()
    )

def update_word_document(docx_file, improvements):
    """Update the Word document with suggested improvements."""
    try:
        doc = Document(docx_file)
        
        # Add improvements to the document
        if improvements:
            # Find sections in the document
            experience_section = None
            skills_section = None
            projects_section = None
            
            # Search for section headings
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.lower()
                if "experience" in text or "work" in text:
                    experience_section = i
                elif "skill" in text:
                    skills_section = i
                elif "project" in text:
                    projects_section = i
            
            # Add experience improvements
            if "Experience" in improvements and experience_section is not None:
                # Find the next section after experience
                next_section = len(doc.paragraphs)
                for i in range(experience_section + 1, len(doc.paragraphs)):
                    if doc.paragraphs[i].text.strip() and doc.paragraphs[i].style.name.startswith('Heading'):
                        next_section = i
                        break
                
                # Add improvements before the next section
                for item in improvements["Experience"]:
                    p = doc.add_paragraph("• " + item, style='List Bullet')
                    # Move paragraph to the right position
                    doc._body._body.insert(next_section, p._p)
            
            # Add skills improvements
            if "Skills" in improvements and skills_section is not None:
                # Find the next section after skills
                next_section = len(doc.paragraphs)
                for i in range(skills_section + 1, len(doc.paragraphs)):
                    if doc.paragraphs[i].text.strip() and doc.paragraphs[i].style.name.startswith('Heading'):
                        next_section = i
                        break
                
                # Combine skills into a single paragraph
                if improvements["Skills"]:
                    skills_text = ", ".join(improvements["Skills"])
                    p = doc.add_paragraph(skills_text)
                    # Move paragraph to the right position
                    doc._body._body.insert(next_section, p._p)
            
            # Add projects improvements
            if "Projects" in improvements and projects_section is not None:
                # Find the next section after projects
                next_section = len(doc.paragraphs)
                for i in range(projects_section + 1, len(doc.paragraphs)):
                    if doc.paragraphs[i].text.strip() and doc.paragraphs[i].style.name.startswith('Heading'):
                        next_section = i
                        break
                
                # Add improvements before the next section
                for item in improvements["Projects"]:
                    p = doc.add_paragraph("• " + item, style='List Bullet')
                    # Move paragraph to the right position
                    doc._body._body.insert(next_section, p._p)
        
        # Save the updated document
        output_bytes = io.BytesIO()
        doc.save(output_bytes)
        output_bytes.seek(0)
        
        return output_bytes.getvalue()
        
    except Exception as e:
        raise Exception(f"Error updating Word document: {str(e)}")

def convert_docx_to_pdf(docx_bytes):
    """Convert a Word document to PDF and save it to a specific location."""
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_bytes.getvalue() if isinstance(docx_bytes, io.BytesIO) else docx_bytes)
            temp_docx_path = temp_docx.name
            
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        # Convert DOCX to PDF
        convert(temp_docx_path, temp_pdf_path)
        
        # Read the resulting PDF
        with open(temp_pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Save the PDF to a specific location
        output_dir = "saved_resumes"  # You can change this to your desired directory
        os.makedirs(output_dir, exist_ok=True)  # Create the directory if it doesn't exist
        
        # Generate a unique filename using timestamp
        import datetime
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"resume_{timestamp}.pdf"
        output_path = os.path.join(output_dir, output_filename)
        
        # Save the PDF
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
        
        print(f"PDF saved to: {os.path.abspath(output_path)}")
            
        # Clean up temporary files
        os.remove(temp_docx_path)
        os.remove(temp_pdf_path)
        
        return pdf_bytes
        
    except Exception as e:
        raise Exception(f"Error converting Word document to PDF: {str(e)}")